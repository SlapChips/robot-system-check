#!/bin/python
from docx import Document
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from pprint import pprint  # Import the pprint function
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
# from docx.enum.text import WD_LIST_NUMBER
import re
import sys
"""WORKING DOC"""


def log_output_to_file(file_name):
    def decorator(func):
        def wrapper(*args, **kwargs):
            original_stdout = sys.stdout
            original_stderr = sys.stderr

            with open(file_name, 'w') as log_file:
                sys.stdout = log_file
                sys.stderr = log_file
                try:
                    result = func(*args, **kwargs)
                finally:
                    sys.stdout = original_stdout
                    sys.stderr = original_stderr

            return result

        return wrapper

    return decorator


def parse_output_xml(file_path):
    """
    Parse the output.xml file defined in the input.

    Return: output_xml
    """
    try:
        # Parse the XML file
        tree = ET.parse(file_path)
        output_xml = tree.getroot()
    except ET.ParseError as e:
        raise ValueError(f"Failed to parse the XML file: {e}")
    # Initialize a dictionary to store the attributes
    return output_xml


def clean_doc_text(doc_text):
    lines = doc_text.split('\n')
    cleaned_lines = []
    for i, line in enumerate(lines):
        current_line = line.strip()  # Remove leading/trailing whitespace
        if i < len(lines) - 1:
            # Remove leading/trailing whitespace from the next line
            next_line = lines[i + 1].strip()
        else:
            next_line = ""
        if current_line and \
            not re.match(r'^\s*[-*]', current_line) and \
                not next_line.startswith(('*', '-')):
            current_line = current_line.lstrip()  # Remove leading whitespace
        cleaned_lines.append(current_line)
    return '\n'.join(cleaned_lines)


def add_caption(doc, caption_type, caption_text):
    """
    Creates a caption using the Cisco Caption style

    Inputs:
    doc :   the document object
    type:   Table, Code, Figure
    Caption:    Caption Text to append
    """
    # Define the Valid Caption Types (Table, Figure, or Code)
    if caption_type.title() in ["Figure", "Code", "Table"]:
        pass
    else:
        raise ValueError(
            "Invalid type specified. Use 'figure', 'code', or 'table'.")
    paragraph = doc.add_paragraph(f'{caption_type.title()}-', style='Caption')
    add_caption_field_code(paragraph=paragraph,
                           caption_type=caption_type.title())
    run = paragraph.add_run(f' {caption_text}')


def add_caption_field_code(paragraph, caption_type):
    """
    invoke this be creating a new paragraph and passing it into this function

    e.g.
        paragraph = doc.add_paragraph(
            f'{caption_type.title()} : {caption_text}',
            style='Caption')
        add_caption_field_code(paragraph=paragraph)
    """
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = f' SEQ {caption_type.title()} \\* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def get_section_summary_results(doc, data, caption):
    # Add an empty line before the table
    add_caption(doc, 'Table', caption)
    # Add an empty line after the caption
    doc.add_paragraph()
    # Add a table with one row and the number of columns equal to the length of
    # the input dictionary
    table = doc.add_table(rows=1, cols=len(data))
    # Apply the custom table style to the table
    table.style = 'Cisco CX Table | Default'
    # Add the header row with labels from the input dictionary
    header_row = table.rows[0].cells
    for idx, label in enumerate(data.keys()):
        header_row[idx].text = label
    # Add the values from the input dictionary as a new row
    values_row = table.add_row().cells
    for idx, value in enumerate(data.values()):
        values_row[idx].text = value
    # Add an empty line after the table
    doc.add_paragraph()


def get_section_statistics(output_xml, section_id):
    # Define the XPath for the specified section
    # (total, tag, or suite) with the given section_id
    xpath = f".//suite/stat[@id='{section_id}']"
    # print(section_id)
    # Find the element that matches the section_id
    section_element = output_xml.find(xpath)

    if section_element is None:
        raise ValueError(
            f"Section with ID '{section_id}' not found in the XML.")

    # Extract the pass, fail, and skip values
    pass_value = section_element.get("pass")
    fail_value = section_element.get("fail")
    skip_value = section_element.get("skip")

    # Create a dictionary with the extracted values
    result_dict = {
        "pass": pass_value,
        "fail": fail_value,
        "skip": skip_value,
    }

    return result_dict
    """
    Returns:
    {'All Tests': {'pass': '20', 'fail': '0', 'skip': '0'}}
    """
    return section_dict


def get_test_sections(output_xml):
    """
    Get the Suite Name and ID Mapping
    <suite id="s1" name="Tests" source="/home/ubutt/robot-dev/tests">
        <suite id="s1-s1" name="Ncs Env" source="/tests/ncs_env.robot">

    Returns Dict:
    {
        's1': {'name': 'Tests', 'source': '/tests'},
        's1-s1': {'name': 'Ncs Env', 'source': '/tests/ncs_env.robot'},
        's1-s2': {'name': 'Security', 'source': '/tests/security.robot'},
        's1-s3': {'name': 'System', 'source': '/tests/system.robot'}}

    """
    # Initialize a dictionary to store the attributes
    suite_attributes = {}
    for suite_element in output_xml.findall(".//suite"):
        if suite_element.get("id") is None:
            continue  # Skip this iteration and move to the next element
        id_value = suite_element.get("id")
        name_value = suite_element.get("name")
        source_value = suite_element.get("source")
        doc_element = suite_element.find("doc")
        if doc_element is not None:
            doc_value = clean_doc(doc_element.text)
        else:
            doc_value = ""
        # Insert the attributes into the dictionary
        # with the suite name as the key
        suite_attributes[id_value] = {
            # "id": id_value,
            "name": name_value,
            "source": source_value,
            "doc": doc_value
        }
    return suite_attributes


def get_tests(output_xml):
    test_list = []
    for test_element in output_xml.findall(".//test"):
        id_value = test_element.get("id")
        name_value = test_element.get("name")
        doc_value = test_element.find("doc").text
        section_value = id_value.split("-t")[0]
        status_value = test_element.find("status").get("status")
        # Insert the attributes into a dictionary
        msg_elements = test_element.findall(".//msg")
        # Iterating to see if there are any kw matching
        steps_list = []

        for kw_element in test_element.findall("kw"):
            kw_name = kw_element.get("name")  # Get the name attribute
            print(kw_name)
            if ('Step.' in kw_name):
                print(f'Found Step > {kw_name}')
                steps_list.append(kw_name)
        print(steps_list)
        # merge the messages output into one value
        # currently each message is seperated by CR
        # need to trim whitespace in each message
        # for msg_element in msg_elements:
        #     msg_element.text = re.sub(r'\s{2,}', ' ', msg_element.text)
        message_value = "\n".join(
            msg_element.text for msg_element in msg_elements)
        message_value = message_value
        test_data = {
            "id":   id_value,
            "name": name_value,
            "doc": doc_value,
            "section": section_value,
            "status": status_value,
            "messages": message_value,
            "procedure": steps_list
        }
        test_list.append(test_data)
    return test_list


def add_vertical_testcase_table(doc, data_dict):
    add_caption(doc, caption_type='table', caption_text=data_dict['name'])
    page_width = doc.sections[0].page_width
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Cisco CX Table | Default'
    # Set the width for the first column
    table.columns[0].width = int(page_width * 0.2)
    table.columns[1].width = int(page_width * 0.8)
    # Specifiy the order of the test case rows, in the table
    key_order = ['section', 'doc', 'procedure', 'status', 'messages']
    for key in key_order:
        # Get the value for the current key, default to empty string if key not found
        value = data_dict.get(key, '')
        row = table.add_row().cells
        row[0].text = key
        if key == 'procedure' and isinstance(value, list):
            # If the key is 'procedure' and the value is a list, add list items as bullets
            print(value)
            value_stripped = [item.replace("Step. ", "") for item in value]
            cell = row[1]
            for item in value_stripped:
                p = cell.add_paragraph(item, style='Step Un-numbered')
        elif key == 'doc':
            """ Reformat the <doc> contents to remove page breaks"""
            value = clean_doc_text(value)
            row[0].text = 'Purpose'
            row[1].text = value
        elif key == 'messages':
            # Applying the Code style to message log
            cell = row[1]
            p = cell.add_paragraph(value, style='Code')
        else:
            # If the key is not 'procedure' or the value is not a list, set the text directly
            row[1].text = value


def clean_doc(value):
    # Remove all carriage returns and extra whitespace
    value = re.sub(r'[\r\n]+', ' ', value)
    value = re.sub(r'\s{2,}', ' ', value)
    # if value has - or * with whitespace on both sides, add a new line:
    value = re.sub(r' - ', '\n - ', value)
    value = re.sub(r'\* ', '\n * ', value)
    return value


def add_testcase_table_in_position(doc,
                                   test_data_dict,
                                   positon='before'
                                   ):
    page_width = doc.sections[0].page_width
    test_section = test_data_dict['section']
    test_id = test_data_dict['id']
    test_name = test_data_dict['name']
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Cisco CX Table | Default'
    # Set the width for the first column
    table.columns[0].width = int(page_width * 0.2)
    table.columns[1].width = int(page_width * 0.8)
    key_order = ['name', 'id', 'section', 'doc',
                 'procedure', 'status', 'messages']
    # Add content to the table cells
    for key in key_order:
        value = test_data_dict.get(key, '')
        row = table.add_row().cells
        row[0].text = key
        if isinstance(value, list):
            # If the value is a list, add list items as bullets
            print(value)
            value_stripped = [item.replace("Step. ", "") for item in value]
            cell = row[1]
            # use existing cell paragraph
            p = cell.paragraphs[0]
            for item in value_stripped:
                p.add_run("• " + item + "\n")
        else:
            row[1].text = value
    pprint(test_data_dict)
    testcase_anchor = create_inline_paragraph(
        doc,
        placeholder=test_section,
        position=positon,
        style_name='Caption'
    )
    # This line assignes the tables to the placeholder location
    # Need to modify Caption code...
    # testcase_anchor.add_run(test_id)
    caption_text = test_id + ' ' + test_name
    add_caption_new(testcase_anchor, 'Table', caption_text)
    testcase_anchor._p.addnext(table._tbl)


def add_caption_new(paragraph, caption_type, caption_text):
    """
    Creates a caption using the Cisco Caption style

    Inputs:
    doc :   the document object
    type:   Table, Code, Figure
    Caption:    Caption Text to append
    """
    # Define the Valid Caption Types (Table, Figure, or Code)
    if caption_type.title() in ["Figure", "Code", "Table"]:
        pass
    else:
        raise ValueError(
            "Invalid type specified. Use 'figure', 'code', or 'table'.")
    caption = paragraph.add_run(f'{caption_type.title()}: {caption_text}')
    # caption = paragraph.add_paragraph(f'{caption_type.title()}: caption_text', style='Caption')
    add_caption_field_code(paragraph=paragraph,
                           caption_type=caption_type.title())
    # run = paragraph.add_run(f' {caption_text}')


def add_content_inline(doc,
                       section_name,
                       placeholder,
                       position='before',
                       style='Normal',
                       ):
    """ Adds a section heading before the placeholder"""
    testcase_anchor = create_inline_paragraph(
        doc,
        placeholder=placeholder,
        position=position,
        style_name=style
    )
    testcase_anchor.add_run(section_name)
    return testcase_anchor


def add_all_test_sections(doc, sections):
    added_sections = {}
    for section_id, section_data in sections.items():
        section_doc = section_data['doc']
        section_name = section_data['name']      
        # evaluate if i need to add a new section, if true create heading and
        # dump section documentation from the test suite docs
        if section_id not in added_sections:
            # We want all section headings and content above the 
            # placeholder text <test_section_placeholder>
            placeholder = 'test_section_placeholder'
            position = 'before'
            # this is a new test suite thats has no heading, adding a heading:
            add_section = add_content_inline(doc,
                                             section_name,
                                             placeholder,
                                             position=position,
                                             style='Heading 2',
                                             )
            add_section_doc = add_content_inline(doc,
                                                 section_doc,
                                                 placeholder,
                                                 position=position,
                                                 style='Normal'
                                                 )
            new_text = f'\n<{section_id}>'
            add_section_placeholder = add_content_inline(doc,
                                                         new_text,
                                                         placeholder,
                                                         position=position,
                                                         style='Normal'
                                                         )
            added_sections[section_id] = True


def create_test_docx(**kwargs):
    """
    Bring everything together and create the docx content

    Defaults:
    table_style_name = 'Cisco CX Table | Default'
    test_output_xml = './output.xml'
    docx_template = 'Template.docx'

    """
    output_docx = kwargs.get('output_docx', 'output.docx')
    test_output_xml = kwargs.get('test_output_xml', './output.xml')
    docx_template = kwargs.get('docx_template', 'Template.docx')
    table_style_name = 'Cisco CX Table | Default'

    # Calling functions top extract test data:
    output_xml = parse_output_xml(test_output_xml)
    sections = get_test_sections(output_xml=output_xml)
    test_list = get_tests(output_xml=output_xml)
    # Create a new Word document
    doc = Document(docx_template)
    table_style_name = 'Cisco CX Table | Default'
    # Check if the style exists in the document's styles
    if table_style_name in doc.styles:
        # Get the style object
        custom_table_style = doc.styles[table_style_name]
    else:
        # If the style doesn't exist, you can create a new style based on it
        custom_table_style = doc.styles.add_style(
            table_style_name, 'Table Normal')

    table = doc.tables[0]
    # Asisgn the Cisco CX Default Style
    table.style = table_style_name
    test_case_tables = []
    add_all_section = add_all_test_sections(doc, sections)
    
    for test_data in test_list:
        test_id = test_data['id']
        section_id = test_data['section']
        test_case = add_testcase_table_in_position(doc, test_data)

    #     # This messed up the whole layout:
    #     # add_vertical_testcase_table(doc, test_data)

    #     """
    #     Content below continue is old code that works...
    #     skipping to test content inline placement
    #     """
    #     continue
    #     section_id = test_data['section']
    #     if section_id not in sections:
    #         continue  # Skip tests without a matching section
    #     # Add a section break if it's a new section
    #     if section_id not in added_sections:
    #         if not doc.sections or doc.sections[-1].footer is None:
    #             section = doc.sections[-1]
    #             section.start_type = WD_SECTION_START.CONTINUOUS
    #             section.start_param = WD_SECTION.NEW_COLUMN
    #             section.footer.is_linked_to_previous = False
    #         # Add a heading for the section
    #         # (using the section name as the heading text)
    #         section_name = sections[section_id]['name']
    #         doc.add_heading(section_name, level=2)
    #         doc.add_paragraph(sections[section_id]['doc'])
    #         added_sections[section_id] = True
    #         section_stats = get_section_statistics(output_xml, section_id)
    #         get_section_summary_results(
    #             doc, section_stats,
    #             caption=f'{section_name} Test Results Summary')
    #     # pprint(test_data)
    #     # test_case_tables.append()
    #     # add_vertical_testcase_table(doc, test_data)

    doc.save(output_docx)
    print(f'Created Output file : {output_docx}')


def find_placeholder(doc, placeholder):
    """
    Find the location of a placeholder, we need this for inline edits of a docx
    template. The input is a seach string that the code will use to iterate over 
    the raw document text.

    Return: Paragraph Number (int)
    This can be used to construct a new parapgraph statement like this:

    insert_before_here = doc.paragraphs[106]

    This creates a new paragraph before the paragraph 106
    """
    # create a list of all paragraphs, painful but neccesary:
    para_list = doc.paragraphs
    # set counter to store location:
    count = 0
    for p in para_list:
        if (placeholder in p.text):
            # found you, you retartd
            p_id = count
        count += 1
    return p_id  # return the paragraph id (int)


def create_inline_paragraph(doc, placeholder, position, style_name=None):
    """
    Create a new paragraph object before the placeholder text:

    placeholder = some text placeholder in docx template
    position = relative position : before / after

    """
    p_id = find_placeholder(doc, placeholder)
    if position.lower() == 'before':
        p = doc.paragraphs[p_id].insert_paragraph_before()
    elif position.lower() == 'after':
        p = doc.paragraphs[p_id]
    if style_name:
        p.style = style_name
    return p  # returning docx paragraph object


if __name__ == "__main__":
    # set ROOT_PATH to absolutio location of code
    root_path = '/Users/ubutt/git/robot-system-check/rhel-dev/'
    test_output_xml = f'{root_path}results/output.xml'
    docx_template = f'{root_path}create_docx/vf-atp-template.docx'
    create_test_docx(
        test_output_xml=test_output_xml,
        docx_template=docx_template
    )
"""

table_style_name = 'Cisco CX Table | Default'
test_output_xml = './output.xml'
docx_template = '/Users/ubutt/git/robot-system-check/docx/update_docx_props/vf-atp-template.docx'
# Calling functions top extract test data:
output_xml = parse_output_xml(test_output_xml)
sections = get_test_sections(output_xml=output_xml)
tests = get_tests(output_xml=output_xml)

Manual Tests:
from lxml import etree as ET
from extract_test import *
test_output_xml = '/Users/ubutt/git/robot-system-check/override_doc/output.xml'
output_xml = parse_output_xml(test_output_xml)
tests = get_tests(output_xml=output_xml)
xml_string = ET.tostring(output_xml, encoding='unicode', method='xml')
"""
