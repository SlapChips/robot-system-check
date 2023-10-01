from docx import Document
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn


def parse_statistics_from_file(file_path, section):
    try:
        # Parse the XML file
        tree = ET.parse(file_path)
        root = tree.getroot()
    except ET.ParseError as e:
        raise ValueError(f"Failed to parse the XML file: {e}")

    # Initialize a dictionary to store the results
    result_dict = {}

    # Define the XPath for the specified section (total, tag, or suite)
    if section == "total":
        xpath = ".//total/stat"
    elif section == "tag":
        xpath = ".//tag/stat"
    elif section == "suite":
        xpath = ".//suite/stat"
    else:
        raise ValueError("Invalid section specified. Use 'total', 'tag', or 'suite'.")

    # Iterate through the selected elements and extract the attributes
    for elem in root.findall(xpath):
        pass_value = elem.get("pass")
        fail_value = elem.get("fail")
        skip_value = elem.get("skip")
        name_value = elem.get("name")

        # If name attribute is not present, use the element's text as name
        if name_value is None:
            name_value = elem.text

        # Create a dictionary entry for the element
        result_dict[name_value] = {
            "pass": pass_value,
            "fail": fail_value,
            "skip": skip_value,
        }
    """
    Returns:
    {'All Tests': {'pass': '20', 'fail': '0', 'skip': '0'}}
    """
    return result_dict


def add_table_to_document_old(doc, data):
    table_style_name = 'Cisco CX Table | Default'
    # Add a table with headers
    table = doc.add_table(rows=1, cols=4)

    # Apply the custom table style to the table
    table.style = table_style_name

    # Set alignment for the "Test Name" cell to center vertically
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add table headers
    table_cells = table.rows[0].cells
    table_cells[0].text = 'Test Name / Tag'
    table_cells[1].text = 'Pass'
    table_cells[2].text = 'Fail'
    table_cells[3].text = 'Skip'

    # Add data to the table
    for test_name, test_data in data.items():
        row = table.add_row().cells
        row[0].text = test_name  # Set the test name in the "Test Name" column
        for i, (label, value) in enumerate(test_data.items()):
            row[i + 1].text = value
    # Add a continuous section break after the table
    doc.add_section(WD_SECTION_START.CONTINUOUS)

def add_table_to_document(doc, data, caption):
    table_style_name = 'Cisco CX Table | Default'
    
    # Add a paragraph for the caption and align it to the left
    table_caption = doc.add_paragraph(style='Caption')
    table_caption.add_run('Table  ')
    field_code = "AUTONUM \* MERGEFORMAT"
    table_caption.add_run().add_field(field_code)
    #table_caption.alignment = WD_ALIGN_VERTICAL.LEFT
    table_caption.add_run(caption)

    
    # Add a table with headers
    table = doc.add_table(rows=1, cols=4)

    # Apply the custom table style to the table
    table.style = table_style_name

    # Set alignment for the "Test Name" cell to center vertically
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add table headers
    table_cells = table.rows[0].cells
    table_cells[0].text = 'Test Name / Tag'
    table_cells[1].text = 'Pass'
    table_cells[2].text = 'Fail'
    table_cells[3].text = 'Skip'

    # Add data to the table
    for test_name, test_data in data.items():
        row = table.add_row().cells
        row[0].text = test_name  # Set the test name in the "Test Name" column
        for i, (label, value) in enumerate(test_data.items()):
            row[i + 1].text = value

    # Add a continuous section break after the table
    doc.add_section(WD_SECTION_START.CONTINUOUS)


def generate_stats(doc):
    stats_collection = ['total', 'tag', 'suite']

    for stats in stats_collection:
        data = parse_statistics_from_file("output.xml", stats)
        if stats == 'total':
            caption = 'Statistics for all tests'
        elif stats == 'tag':
            caption == 'Statistics by test TAG'
        elif stats == 'suite':
            caption == 'Statistics organised by Robot File executed'
        add_table_to_document(doc, data, caption)
        # Save the document


doc = Document('Template.docx')  # Create a new document for each section
generate_stats(doc)
doc.save('test_results_with_custom_style.docx')
