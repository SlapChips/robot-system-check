import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL


def parse_statistics_from_file(file_path, section):
    try:
        # Parse the XML file
        tree = ET.parse(file_path)
        root = tree.getroot()
    except ET.ParseError as e:
        raise ValueError(f"Failed to parse the XML file: {e}")

    # Initialize a dictionary to store the results
    result_dict = {}

    # Define the XPath for the specified section (total, tag, or suite)
    if section == "total":
        xpath = ".//total/stat"
    elif section == "tag":
        xpath = ".//tag/stat"
    elif section == "suite":
        xpath = ".//suite/stat"
    else:
        raise ValueError(
            "Invalid section specified. Use 'total', 'tag', or 'suite'.")

    # Iterate through the selected elements and extract the attributes
    for elem in root.findall(xpath):
        pass_value = elem.get("pass")
        fail_value = elem.get("fail")
        skip_value = elem.get("skip")
        name_value = elem.get("name")

        # If name attribute is not present, use the element's text as name
        if name_value is None:
            name_value = elem.text

        # Create a dictionary entry for the element
        result_dict[name_value] = {
            "pass": pass_value,
            "fail": fail_value,
            "skip": skip_value,
        }
    """
    Returns:
    {'All Tests': {'pass': '20', 'fail': '0', 'skip': '0'}}
    """
    return result_dict


# Example usage:
file_path = "output.xml"  # Replace with the path to your XML file
section = "suite"  # Specify the desired section ("total", "tag", or "suite")

results = parse_statistics_from_file(file_path, section)
print(results)


# Sample data in the format you provided
data = results

# Open the Word document containing the table style
custom_style_doc = Document('Template.docx')

# Get the table style name (e.g., "Cisco CX Table | Default")
table_style_name = 'Cisco CX Table | Default'

# Check if the style exists in the document's styles
if table_style_name in custom_style_doc.styles:
    # Get the style object
    custom_table_style = custom_style_doc.styles[table_style_name]
else:
    # If the style doesn't exist, you can create a new style based on it
    custom_table_style = custom_style_doc.styles.add_style(table_style_name,
                                                           'Table Normal')

# Create a new Word document
doc = Document("Template.docx")

# Add a table with headers
table = doc.add_table(rows=1, cols=4)

# Apply the custom table style to the table
table.style = custom_table_style

# Set alignment for the "Test Name" cell to center vertically and horizontally
cell = table.cell(0, 0)
cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
# cell.paragraphs[0].alignment = WD_ALIGN_HORIZONTAL.CENTER

# Add table headers
table_cells = table.rows[0].cells
table_cells[0].text = 'Test Name'
table_cells[1].text = 'Pass'
table_cells[2].text = 'Fail'
table_cells[3].text = 'Skip'

# Add data to the table
for test_name, test_data in data.items():
    row = table.add_row().cells
    row[0].text = test_name  # Set the test name in the "Test Name" column
    for i, (label, value) in enumerate(test_data.items()):
        row[i + 1].text = value

# Save the document
doc.save('test_results_with_custom_style.docx')
