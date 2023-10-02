#!/bin/python
from docx import Document
import xml.etree.ElementTree as ET
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.oxml.shared import OxmlElement, qn
from pprint import pprint  # Import the pprint function
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
import sys
"""WORKING DOC"""


def log_output_to_file(file_name):
    def decorator(func):
        def wrapper(*args, **kwargs):
            original_stdout = sys.stdout
            original_stderr = sys.stderr

            with open(file_name, 'w') as log_file:
                sys.stdout = log_file
                sys.stderr = log_file
                try:
                    result = func(*args, **kwargs)
                finally:
                    sys.stdout = original_stdout
                    sys.stderr = original_stderr

            return result

        return wrapper

    return decorator


def parse_output_xml(file_path):
    """
    Parse the output.xml file defined in the input.

    Return: output_xml
    """
    try:
        # Parse the XML file
        tree = ET.parse(file_path)
        output_xml = tree.getroot()
    except ET.ParseError as e:
        raise ValueError(f"Failed to parse the XML file: {e}")
    # Initialize a dictionary to store the attributes
    return output_xml


def clean_doc_text(doc_text):
    lines = doc_text.split('\n')
    cleaned_lines = []
    for i, line in enumerate(lines):
        current_line = line.strip()  # Remove leading/trailing whitespace
        if i < len(lines) - 1:
            next_line = lines[i + 1].strip()  # Remove leading/trailing whitespace from the next line
        else:
            next_line = ""
        if current_line and not re.match(r'^\s*[-*]', current_line) and not next_line.startswith(('*', '-')):
            current_line = current_line.lstrip()  # Remove leading whitespace
        cleaned_lines.append(current_line)
    return '\n'.join(cleaned_lines)


def add_caption(doc, caption_type, caption_text):
    """
    Creates a caption using the Cisco Caption style

    Inputs:
    doc :   the document object
    type:   Table, Code, Figure
    Caption:    Caption Text to append
    """
    # Define the Valid Caption Types (Table, Figure, or Code)
    if caption_type.title() in ["Figure", "Code", "Table"]:
        pass
    else:
        raise ValueError("Invalid type specified. Use 'figure', 'code', or 'table'.")
 
    paragraph = doc.add_paragraph(f'{caption_type.title()}-', style='Caption')
    add_caption_field_code(paragraph=paragraph, caption_type=caption_type.title())
    run = paragraph.add_run(f' {caption_text}')


def add_caption_field_code(paragraph, caption_type):
    """
    invoke this be creating a new paragraph and passing it into this function

    e.g.
        paragraph = doc.add_paragraph(f'{caption_type.title()} : {caption_text}', style='Caption')
        add_caption_field_code(paragraph=paragraph)
    """
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = f' SEQ {caption_type.title()} \\* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def get_section_summary_results(doc, data, caption):
    # Add an empty line before the table
    add_caption(doc, 'Table', caption)
    # Add an empty line after the caption
    doc.add_paragraph()
    # Add a table with one row and the number of columns equal to the length of the input dictionary
    table = doc.add_table(rows=1, cols=len(data))
    # Apply the custom table style to the table
    table.style = 'Cisco CX Table | Default'
    # Add the header row with labels from the input dictionary
    header_row = table.rows[0].cells
    for idx, label in enumerate(data.keys()):
        header_row[idx].text = label
    # Add the values from the input dictionary as a new row
    values_row = table.add_row().cells
    for idx, value in enumerate(data.values()):
        values_row[idx].text = value
    # Add an empty line after the table
    doc.add_paragraph()


def create_table_in_docx(doc, data):
    # Determine the number of rows and columns based on the data
    num_rows = len(data)
    num_columns = len(next(iter(data.values()))) + 1  # Add 1 for the "Test Name" column
    # print(f'Creating {num_rows} rows, and {num_columns} colums')
    # Add a table with dynamic rows and columns
    table = doc.add_table(rows=num_rows, cols=num_columns)

    # Apply the custom table style to the table
    table.style = 'Cisco CX Table | Default'

    # Set alignment for the "Test Name" cell to center vertically
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add table headers
    table_cells = table.rows[0].cells
    table_cells[0].text = 'Test Name'

    # Add headers for additional columns (if any)
    for i, label in enumerate(data.values()):
        table_cells[i + 1].text = label

    # Add data to the table
    for row_index, (test_name, test_data) in enumerate(data.items(), start=1):
        row = table.rows[row_index].cells
        row[0].text = test_name  # Set the test name in the "Test Name" column
        for i, value in enumerate(test_data.values(), start=1):
            row[i].text = value


def get_section_statistics(output_xml, section_id):
    # Define the XPath for the specified section (total, tag, or suite) with the given section_id
    xpath = f".//suite/stat[@id='{section_id}']"
    # print(section_id)
    # Find the element that matches the section_id
    section_element = output_xml.find(xpath)

    if section_element is None:
        raise ValueError(f"Section with ID '{section_id}' not found in the XML.")

    # Extract the pass, fail, and skip values
    pass_value = section_element.get("pass")
    fail_value = section_element.get("fail")
    skip_value = section_element.get("skip")

    # Create a dictionary with the extracted values
    result_dict = {
        "pass": pass_value,
        "fail": fail_value,
        "skip": skip_value,
    }

    return result_dict
    """
    Returns:
    {'All Tests': {'pass': '20', 'fail': '0', 'skip': '0'}}
    """
    return section_dict


def get_tests_from_output_xml(file_path, section):
    try:
        # Parse the XML file
        tree = ET.parse(file_path)
        root = tree.getroot()
    except ET.ParseError as e:
        raise ValueError(f"Failed to parse the XML file: {e}")

    # Initialize a dictionary to store the results
    result_dict = {}

    # Define the XPath for the specified section (total, tag, or suite)
    if section == "total":
        xpath = ".//total/stat"
    elif section == "tag":
        xpath = ".//tag/stat"
    elif section == "suite":
        xpath = ".//suite/stat"
    else:
        raise ValueError("Invalid section specified. Use 'total', 'tag', or 'suite'.")

    # Iterate through the selected elements and extract the attributes
    for elem in root.findall(xpath):
        pass_value = elem.get("pass")
        fail_value = elem.get("fail")
        skip_value = elem.get("skip")
        name_value = elem.get("name")

        # If name attribute is not present, use the element's text as name
        if name_value is None:
            name_value = elem.text

        # Create a dictionary entry for the element
        result_dict[name_value] = {
            "pass": pass_value,
            "fail": fail_value,
            "skip": skip_value,
        }
    """
    Returns:
    {'All Tests': {'pass': '20', 'fail': '0', 'skip': '0'}}
    """
    return result_dict


def get_test_sections(output_xml):
    """
    Get the Suite Name and ID Mapping
    <suite id="s1" name="Tests" source="/home/ubutt/robot-dev/tests">
        <suite id="s1-s1" name="Ncs Env" source="/home/ubutt/robot-dev/tests/ncs_env.robot">

    Returns Dict:
    {
        's1': {'name': 'Tests', 'source': '/home/ubutt/robot-dev/tests'}, 
        's1-s1': {'name': 'Ncs Env', 'source': '/home/ubutt/robot-dev/tests/ncs_env.robot'}, 
        's1-s2': {'name': 'Security', 'source': '/home/ubutt/robot-dev/tests/security.robot'}, 
        's1-s3': {'name': 'System', 'source': '/home/ubutt/robot-dev/tests/system.robot'}}

    """
    # Initialize a dictionary to store the attributes
    suite_attributes = {}
    for suite_element in output_xml.findall(".//suite"):
        if suite_element.get("id") == None:
            continue  # Skip this iteration and move to the next element
        id_value = suite_element.get("id")
        name_value = suite_element.get("name")
        source_value = suite_element.get("source")
        doc_element = suite_element.find("doc")
        doc_value = clean_doc_text(doc_element.text) if doc_element is not None else ""
        # Insert the attributes into the dictionary with the suite name as the key
        suite_attributes[id_value] = {
            # "id": id_value,
            "name": name_value,
            "source": source_value,
            "doc": doc_value
        }
    return suite_attributes


def get_tests(output_xml):
    test_list = []
    for test_element in output_xml.findall(".//test"):
        id_value = test_element.get("id")
        name_value = test_element.get("name")
        doc_value = test_element.find("doc").text
        section_value = id_value.split("-t")[0]  # Extract the section ID without
        status_value = test_element.find("status").get("status")
        # Insert the attributes into a dictionary
        msg_elements = test_element.findall(".//msg")
        message_value = "\n".join(msg_element.text for msg_element in msg_elements)
        test_data = {
            "name": name_value,
            "doc": doc_value,
            "section": section_value,
            "status": status_value,
            "messages": message_value
        }
        test_list.append(test_data)
    return test_list


def add_vertical_testcase_table__old(doc, test_data):
    # Create a table with one column and a row for each key-value pair in test_data
    table = doc.add_table(rows=len(test_data), cols=2)
    table.style = 'Cisco CX Table | Default'
    doc.add_paragraph()

    # Set the width of the first column to make it a header column
    table.columns[0].width = Cm(4)  # Adjust the width as needed

    # Add data rows based on the input
    for idx, (header, value) in enumerate(test_data.items()):
        # Add the header in the first column
        header_cell = table.cell(idx, 0)
        header_cell.text = header
        header_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add the value in the second column
        value_cell = table.cell(idx, 1)
        value_cell.text = value

    # Apply formatting to the header column (e.g., bold)
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True


def add_vertical_testcase_table(doc, data_dict):
    add_caption(doc, caption_type='table', caption_text=data_dict['name'])
    page_width = doc.sections[0].page_width
    pprint(page_width)
    pprint(f'first_column_width = {int(page_width * 0.8)}')
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Cisco CX Table | Default'  # You can change the table style as needed
    
    # Set the width for the first column
    table.columns[0].width = int(page_width * 0.2)
    table.columns[1].width = int(page_width * 0.8)

    for key, value in data_dict.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value


def add_testcase_table(doc, test_data):
    # Create a table with headers: Test Name, Test Status, Test Documentation, Test Message Log
    add_caption(doc, "table", test_data['name'])
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Cisco CX Table | Default'
    doc.add_paragraph()

    # pprint(test_data)
    # Add the table headers
    header_row = table.rows[0].cells
    header_row[0].text = 'Test Name'
    header_row[1].text = 'Test Status'
    header_row[2].text = 'Test Documentation'
    header_row[3].text = 'Test Message Log'
    # pprint(test_data.keys())
    # pprint(test_data['name'])
    # Add data rows based on the input
    row = table.add_row().cells
    row[0].text = test_data['name']
    row[1].text = test_data['status']
    row[2].text = test_data['doc']
    row[3].text = test_data['messages']


@log_output_to_file('output_log.txt')
def create_test_docx():
    """
    Bring everything together and create the docx content

    Defaults:
    table_style_name = 'Cisco CX Table | Default'
    test_output_xml = './output.xml'
    docx_template = 'Template.docx'

    """
    table_style_name = 'Cisco CX Table | Default'
    test_output_xml = './output.xml'
    docx_template = 'Template.docx'
    # Calling functions top extract test data:
    output_xml = parse_output_xml(test_output_xml)
    sections = get_test_sections(output_xml=output_xml)
    tests = get_tests(output_xml=output_xml)
    # Create a new Word document
    custom_style_doc = Document('Template.docx')
    table_style_name = 'Cisco CX Table | Default'
    # Check if the style exists in the document's styles
    if table_style_name in custom_style_doc.styles:
        # Get the style object
        custom_table_style = custom_style_doc.styles[table_style_name]
    else:
        # If the style doesn't exist, you can create a new style based on it
        custom_table_style = custom_style_doc.styles.add_style(
            table_style_name, 'Table Normal')  
    doc = Document(docx_template)
    # create a blank table to store the data in the document
    table = doc.tables[0]
    # Asisgn the Cisco CX Default Style
    table.style = table_style_name
    added_sections = {}
    # Iterate over the tests and organize them by section
    for test_data in tests:
        section_id = test_data['section']
        if section_id not in sections:
            continue  # Skip tests without a matching section
        # Add a section break if it's a new section
        if section_id not in added_sections:
            if not doc.sections or doc.sections[-1].footer is None:
                section = doc.sections[-1]
                section.start_type = WD_SECTION_START.CONTINUOUS
                section.start_param = WD_SECTION.NEW_COLUMN
                section.footer.is_linked_to_previous = False
            # Add a heading for the section 
            # (using the section name as the heading text)
            section_name = sections[section_id]['name']
            doc.add_heading(section_name, level=1)
            doc.add_paragraph(sections[section_id]['doc'])
            added_sections[section_id] = True
            section_stats = get_section_statistics(output_xml, section_id)
            # pprint(section_stats)
            get_section_summary_results(
                doc, section_stats,
                caption=f'{section_name} Test Results Summary')
        pprint(test_data)
        add_vertical_testcase_table(doc, test_data)
    doc.add_paragraph()
    # Save the document
    doc.save('test_results.docx')


create_test_docx()
